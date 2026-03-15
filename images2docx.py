from docx import Document
from docx.shared import Inches
from PIL import Image
import sys
import os

MAX_WIDTH_INCHES = 6.5  # Fits standard Word page width


def images_to_docx(image_paths, output_docx, images_per_page=1):
    doc = Document()
    count = 0

    for i, path in enumerate(image_paths):
        filename = os.path.basename(path)
        name_without_ext = os.path.splitext(filename)[0]

        # Add heading
        doc.add_heading(name_without_ext, level=2)

        # Load image to determine scaling
        img = Image.open(path)
        width_px, height_px = img.size
        dpi = img.info.get("dpi", (72, 72))[0]

        width_in = width_px / dpi
        scale = min(1.0, MAX_WIDTH_INCHES / width_in)

        doc.add_picture(path, width=Inches(width_in * scale))

        count += 1

        # Add spacing
        doc.add_paragraph()

        # Insert page break if page is full
        if count >= images_per_page and i != len(image_paths) - 1:
            doc.add_page_break()
            count = 0

    doc.save(output_docx)
    print(f"Saved {len(image_paths)} images to {output_docx}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python images_to_docx.py output.docx image1.jpg image2.png ... [images_per_page]")
        sys.exit(1)

    output_docx = sys.argv[1]

    # Detect optional images_per_page argument
    try:
        images_per_page = int(sys.argv[-1])
        image_files = sys.argv[2:-1]
    except ValueError:
        images_per_page = 1
        image_files = sys.argv[2:]

    images_to_docx(image_files, output_docx, images_per_page)